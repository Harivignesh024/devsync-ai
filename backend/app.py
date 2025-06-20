import os
import uuid
import requests
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from dotenv import load_dotenv
from docx import Document
from db import init_db, register_user, validate_user, save_summary, get_summaries

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)

# Secret Key
app.config["JWT_SECRET_KEY"] = os.getenv("JWT_SECRET_KEY", "devsync-secret")
jwt = JWTManager(app)

# DB init
init_db()

# HuggingFace API Setup
HUGGINGFACE_API_KEY = os.getenv("HUGGINGFACE_API_KEY")
HUGGINGFACE_API_URL = "https://api-inference.huggingface.co/models/mistralai/Mixtral-8x7B-Instruct-v0.1"
headers = {
    "Authorization": f"Bearer {HUGGINGFACE_API_KEY}"
}

def query_huggingface(prompt):
    response = requests.post(HUGGINGFACE_API_URL, headers=headers, json={"inputs": prompt})
    response.raise_for_status()
    output = response.json()
    return output[0]["generated_text"] if isinstance(output, list) else output

def save_docx(content, filename):
    doc = Document()
    doc.add_heading("DevSync AI - Summary", 0)
    doc.add_paragraph(content)
    os.makedirs("summaries", exist_ok=True)
    filepath = os.path.join("summaries", filename)
    doc.save(filepath)
    return filepath

@app.route("/api/register", methods=["POST"])
def register():
    data = request.get_json()
    username = data.get("username")
    password = data.get("password")
    if register_user(username, password):
        return jsonify({"msg": "User registered"}), 200
    return jsonify({"msg": "Username already exists"}), 409

@app.route("/api/login", methods=["POST"])
def login():
    data = request.get_json()
    username = data.get("username")
    password = data.get("password")
    if validate_user(username, password):
        token = create_access_token(identity=username)
        return jsonify({"access_token": token}), 200
    return jsonify({"msg": "Invalid credentials"}), 401

@app.route("/api/generate", methods=["POST"])
@jwt_required()
def generate_summary():
    data = request.get_json()
    events = data.get("events", [])
    user = get_jwt_identity()

    text = "\n".join(events).strip()
    if not text:
        return jsonify({"summary": "No events provided."})

    prompt = (
        f"You are an AI that writes daily standup summaries for developers. "
        f"Summarize the following updates for {user} into a short, natural-sounding paragraph:\n{text}\n\nSummary:"
    )

    try:
        result = query_huggingface(prompt)
    except Exception as e:
        return jsonify({"summary": f"Error: {str(e)}"}), 500

    cleaned = result.split("Summary:")[-1].strip()
    final_summary = f"{user}:\n{cleaned}"
    save_summary(user, final_summary)

    filename = f"{user.replace(' ', '_')}_summary.docx"
    save_docx(cleaned, filename)

    return jsonify({"summary": final_summary, "doc_url": f"/api/download/{filename}"}), 200

@app.route("/api/team-summary", methods=["POST"])
@jwt_required()
def team_summary():
    data = request.json.get("team", [])

    if not data:
        return jsonify({"error": "No team data provided"}), 400

    prompt = "Generate a professional team summary from the following updates:\n"
    for member in data:
        user = member.get("user", "Anonymous")
        updates = member.get("updates", [])
        if updates:
            for update in updates:
                prompt += f"{user}: {update}\n"

    try:
        result = query_huggingface(prompt)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    summary_text = result.strip().split(":", 1)[-1].strip()
    filename = f"{uuid.uuid4()}.docx"
    save_docx(summary_text, filename)

    return jsonify({
        "summary": summary_text,
        "doc_url": f"/api/download/{filename}"
    })

@app.route("/api/history", methods=["GET"])
@jwt_required()
def get_history():
    user = get_jwt_identity()
    data = get_summaries(user)
    return jsonify({"history": data})

@app.route("/api/github/<username>")
def get_github(username):
    try:
        response = requests.get(f"https://api.github.com/users/{username}")
        return jsonify(response.json())
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/download/<filename>")
def download(filename):
    return send_from_directory("summaries", filename, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)

